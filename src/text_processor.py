import asyncio
from nltk.tokenize import sent_tokenize, word_tokenize  # type: ignore
from nltk.corpus import stopwords  # type: ignore
from nltk.stem import PorterStemmer  # type: ignore
from pathlib import Path
from docx import Document # type: ignore 
from typing import List, Optional
import subprocess
import chardet # type: ignore
import pypdf # type: ignore
import re

from config.logger import logger

stemmer = PorterStemmer()
stop_words = set(stopwords.words('italian'))


def clean_text(text: str) -> str:
    """Text cleaning and normalization"""
    # Remove special characters except basic punctuation
    text = re.sub(r'[^\w\s\.\!\?\,\;\:]', ' ', text)
    # Remove extra whitespaces
    text = ' '.join(text.split())
    return text.strip()


async def extract_sentences(text: str) -> List[str]:
    """Extract sentence tokens from a text"""
    return await asyncio.to_thread(
        lambda: [s.strip() for s in sent_tokenize(text, language="italian") if len(s.strip()) > 10]
    )


def tokenize_for_bm25(text: str) -> List[str]:
    """Tokenization for Italian BM25"""
    text = text.lower()
    tokens = word_tokenize(text, language='italian')
    filtered_tokens = [
        stemmer.stem(token) for token in tokens 
        if token.isalnum() and token not in stop_words
    ]
    return filtered_tokens


def detect_file_encoding(filepath: Path) -> str:
    """Automatically detects the encoding of a text file."""
    try:
        with open(filepath, 'rb') as f:
            result = chardet.detect(f.read())
        return result.get('encoding') or 'utf-8'
    except Exception:
        return 'utf-8'


def read_text_file(filepath: Path, encoding: Optional[str] = None) -> str:
    """Synchronous reading of text files with encoding detection."""
    encoding = encoding or detect_file_encoding(filepath)
    try:
        with open(filepath, 'r', encoding=encoding, errors='ignore') as f:
            return f.read()
    except Exception as e:
        logger.error(f"❌ Error reading text file {filepath}: {e}")
        return ""
    

def read_pdf_file(filepath: Path) -> str:
    """Extracts text from a PDF file."""
    try:
        with open(filepath, "rb") as f:
            reader = pypdf.PdfReader(f)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        logger.error(f"❌ Error reading PDF {filepath}: {e}")
        return ""


def read_docx_file(filepath: Path) -> str:
    """Extracts text from a DOCX file (paragraphs + tables)."""
    try:
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extracting text from tables
        table_texts = [
            " | ".join(cell.text.strip() for cell in row.cells)
            for table in doc.tables for row in table.rows
            if any(cell.text.strip() for cell in row.cells)
        ]

        return "\n".join(paragraphs + table_texts)
    except Exception as e:
        logger.error(f"❌ Error reading DOCX {filepath}: {e}")
        return ""


def read_doc_file(filepath: Path) -> str:
    """Reads legacy DOC files using antiword."""
    try:
        result = subprocess.run(
            ['antiword', str(filepath)],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode == 0:
            return result.stdout
        raise RuntimeError("antiword execution failed")
    except FileNotFoundError:
        logger.warning(f"⚠️ antiword not available for {filepath.name}, using raw text fallback")
        return read_text_file(filepath)
    except Exception as e:
        logger.error(f"❌ Error reading DOC {filepath}: {e}")
        return ""